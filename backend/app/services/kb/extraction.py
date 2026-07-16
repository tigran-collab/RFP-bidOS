"""Text extraction for knowledge-base source documents.

Extracts text while preserving structure so responses can cite the real
source: page boundaries for PDFs, heading sections for DOCX, and sheet
names + cell ranges for spreadsheets. Heavy parsers (pypdf, PyMuPDF,
python-docx, openpyxl, pytesseract) are imported lazily so the app and the
test suite run without every optional dependency installed.

Uploaded document content is untrusted data. ``scan_for_injection`` flags
prompt-injection-style text so downstream drafting can neutralize it; the
extractor never executes or interprets document instructions.
"""

from __future__ import annotations

import csv
import re
from dataclasses import dataclass, field
from pathlib import Path

# A single document is capped so a hostile/huge file cannot exhaust memory.
MAX_TOTAL_CHARS = 4_000_000
MAX_SEGMENT_CHARS = 200_000

PDF_EXTS = {"pdf"}
DOCX_EXTS = {"docx"}
XLSX_EXTS = {"xlsx", "xlsm", "xltx"}
CSV_EXTS = {"csv", "tsv"}
TXT_EXTS = {"txt", "text", "md"}
IMAGE_EXTS = {"png", "jpg", "jpeg", "tif", "tiff", "gif", "bmp", "webp"}

SUPPORTED_EXTS = PDF_EXTS | DOCX_EXTS | XLSX_EXTS | CSV_EXTS | TXT_EXTS | IMAGE_EXTS

# Prompt-injection patterns commonly embedded in adversarial documents. Matches
# are advisory flags for reviewers + drafting guardrails, not hard rejections.
_INJECTION_PATTERNS: tuple[tuple[str, re.Pattern[str]], ...] = (
    ("ignore_instructions", re.compile(r"ignore\s+(all\s+|the\s+|any\s+)?(previous|prior|above|earlier)\s+instructions", re.I)),
    ("disregard", re.compile(r"disregard\s+(all\s+|the\s+|any\s+)?(previous|prior|above|earlier|system)", re.I)),
    ("system_prompt", re.compile(r"\bsystem\s+prompt\b", re.I)),
    ("you_are_now", re.compile(r"you\s+are\s+now\b", re.I)),
    ("new_instructions", re.compile(r"\bnew\s+instructions\b", re.I)),
    ("role_directive", re.compile(r"^\s*(assistant|system)\s*:", re.I | re.M)),
    ("override", re.compile(r"\boverride\s+(the\s+)?(system|previous|prior|rules|instructions)", re.I)),
    ("reveal_prompt", re.compile(r"(reveal|print|show)\s+(your\s+)?(system\s+)?(prompt|instructions)", re.I)),
    ("do_not_cite", re.compile(r"do\s+not\s+(cite|mention|verify|check)", re.I)),
)


@dataclass
class ExtractedSegment:
    text: str
    page_number: int | None = None
    section: str | None = None
    sheet_name: str | None = None
    cell_range: str | None = None


@dataclass
class ExtractionResult:
    file_type: str
    segments: list[ExtractedSegment] = field(default_factory=list)
    page_count: int | None = None
    sheet_names: list[str] = field(default_factory=list)
    truncated: bool = False
    error: str | None = None

    @property
    def text(self) -> str:
        return "\n\n".join(seg.text for seg in self.segments if seg.text)

    @property
    def has_text(self) -> bool:
        return any(seg.text.strip() for seg in self.segments)


def normalize_file_type(path: str | Path, mime_type: str | None = None) -> str:
    ext = Path(path).suffix.lower().lstrip(".")
    if ext:
        return ext
    mime_map = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
        "text/csv": "csv",
        "text/plain": "txt",
        "image/png": "png",
        "image/jpeg": "jpeg",
    }
    return mime_map.get((mime_type or "").split(";", 1)[0].lower(), "")


def is_supported(file_type: str) -> bool:
    return file_type.lower() in SUPPORTED_EXTS


def scan_for_injection(segments: list[ExtractedSegment]) -> list[dict]:
    """Return prompt-injection flags found in the extracted segments."""
    flags: list[dict] = []
    for seg in segments:
        for name, pattern in _INJECTION_PATTERNS:
            match = pattern.search(seg.text)
            if match:
                start = max(0, match.start() - 40)
                end = min(len(seg.text), match.end() + 40)
                flags.append(
                    {
                        "pattern": name,
                        "page": seg.page_number,
                        "sheet": seg.sheet_name,
                        "snippet": seg.text[start:end].strip().replace("\n", " "),
                    }
                )
    return flags


def extract_document(path: str | Path, file_type: str | None = None) -> ExtractionResult:
    path = Path(path)
    ftype = (file_type or normalize_file_type(path)).lower()
    if not path.exists():
        return ExtractionResult(file_type=ftype, error="File not found")
    try:
        if ftype in PDF_EXTS:
            result = _extract_pdf(path)
        elif ftype in DOCX_EXTS:
            result = _extract_docx(path)
        elif ftype in XLSX_EXTS:
            result = _extract_xlsx(path)
        elif ftype in CSV_EXTS:
            result = _extract_csv(path, ftype)
        elif ftype in TXT_EXTS:
            result = _extract_txt(path)
        elif ftype in IMAGE_EXTS:
            result = _extract_image(path, ftype)
        else:
            return ExtractionResult(
                file_type=ftype, error=f"Unsupported file type: {ftype or 'unknown'}"
            )
    except Exception as exc:  # noqa: BLE001 - surface any parser failure cleanly
        return ExtractionResult(file_type=ftype, error=f"Extraction failed: {exc}")

    result = _enforce_caps(result)
    return result


# --- format-specific extractors ----------------------------------------------


def _extract_pdf(path: Path) -> ExtractionResult:
    segments: list[ExtractedSegment] = []
    page_count = 0
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            segments.append(ExtractedSegment(text=text, page_number=index))
    except Exception:
        # Fallback to PyMuPDF when available.
        try:
            import fitz  # type: ignore

            segments = []
            with fitz.open(str(path)) as pdf:
                page_count = pdf.page_count
                for index, page in enumerate(pdf, start=1):
                    segments.append(
                        ExtractedSegment(text=page.get_text(), page_number=index)
                    )
        except Exception as exc:  # noqa: BLE001
            return ExtractionResult(
                file_type="pdf", error=f"PDF parse failed (pypdf and PyMuPDF): {exc}"
            )
    return ExtractionResult(file_type="pdf", segments=segments, page_count=page_count)


def _extract_docx(path: Path) -> ExtractionResult:
    try:
        import docx  # type: ignore
    except Exception as exc:  # noqa: BLE001
        return ExtractionResult(
            file_type="docx",
            error=(
                "python-docx is not installed. Install it to extract DOCX files: "
                f"pip install python-docx ({exc})"
            ),
        )

    document = docx.Document(str(path))
    segments: list[ExtractedSegment] = []
    current_section = "Body"
    buffer: list[str] = []

    def flush() -> None:
        text = "\n".join(line for line in buffer if line is not None).strip()
        if text:
            segments.append(ExtractedSegment(text=text, section=current_section))
        buffer.clear()

    for para in document.paragraphs:
        style = (para.style.name if para.style else "") or ""
        if style.lower().startswith("heading"):
            flush()
            current_section = para.text.strip() or current_section
            buffer.append(para.text)
        else:
            if para.text:
                buffer.append(para.text)
    flush()

    # Tables → flattened text appended as their own section.
    for t_index, table in enumerate(document.tables, start=1):
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows_text.append(" | ".join(cells))
        if rows_text:
            segments.append(
                ExtractedSegment(
                    text="\n".join(rows_text), section=f"Table {t_index}"
                )
            )

    return ExtractionResult(file_type="docx", segments=segments, page_count=None)


def _extract_xlsx(path: Path) -> ExtractionResult:
    try:
        from openpyxl import load_workbook  # type: ignore
    except Exception as exc:  # noqa: BLE001
        return ExtractionResult(
            file_type="xlsx", error=f"openpyxl is not installed: {exc}"
        )

    workbook = load_workbook(str(path), read_only=True, data_only=True)
    segments: list[ExtractedSegment] = []
    sheet_names: list[str] = []
    try:
        for sheet in workbook.worksheets:
            sheet_names.append(sheet.title)
            lines: list[str] = []
            max_row = 0
            max_col = 0
            for row in sheet.iter_rows():
                cells = []
                for cell in row:
                    value = cell.value
                    if value is None or value == "":
                        continue
                    cells.append(f"{cell.coordinate}={value}")
                    max_row = max(max_row, cell.row or 0)
                    max_col = max(max_col, cell.column or 0)
                if cells:
                    lines.append(" ".join(cells))
            if lines:
                cell_range = None
                if max_row and max_col:
                    from openpyxl.utils import get_column_letter  # type: ignore

                    cell_range = f"A1:{get_column_letter(max_col)}{max_row}"
                segments.append(
                    ExtractedSegment(
                        text="\n".join(lines),
                        sheet_name=sheet.title,
                        cell_range=cell_range,
                    )
                )
    finally:
        workbook.close()

    return ExtractionResult(
        file_type="xlsx", segments=segments, sheet_names=sheet_names, page_count=None
    )


def _extract_csv(path: Path, ftype: str) -> ExtractionResult:
    delimiter = "\t" if ftype == "tsv" else ","
    lines: list[str] = []
    with path.open("r", encoding="utf-8", errors="replace", newline="") as handle:
        reader = csv.reader(handle, delimiter=delimiter)
        for row in reader:
            if any(cell.strip() for cell in row):
                lines.append(" | ".join(cell.strip() for cell in row))
    text = "\n".join(lines)
    segments = [ExtractedSegment(text=text, section="CSV")] if text else []
    return ExtractionResult(file_type=ftype, segments=segments, page_count=None)


def _extract_txt(path: Path) -> ExtractionResult:
    text = path.read_text(encoding="utf-8", errors="replace")
    segments = [ExtractedSegment(text=text)] if text.strip() else []
    return ExtractionResult(file_type="txt", segments=segments, page_count=None)


def _extract_image(path: Path, ftype: str) -> ExtractionResult:
    """OCR an image when pytesseract + Pillow + the tesseract binary are all
    available. Otherwise return no text with a clear, non-fatal note — matching
    the app's existing "no OCR by default" stance while leaving the hook."""
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except Exception:
        return ExtractionResult(
            file_type=ftype,
            error=(
                "OCR is not available (install pytesseract + Pillow and the "
                "tesseract binary to extract text from images)."
            ),
        )
    try:
        text = pytesseract.image_to_string(Image.open(str(path)))
    except Exception as exc:  # noqa: BLE001
        return ExtractionResult(file_type=ftype, error=f"OCR failed: {exc}")
    segments = [ExtractedSegment(text=text, section="OCR")] if text.strip() else []
    return ExtractionResult(file_type=ftype, segments=segments, page_count=None)


def _enforce_caps(result: ExtractionResult) -> ExtractionResult:
    total = 0
    capped: list[ExtractedSegment] = []
    for seg in result.segments:
        text = seg.text or ""
        if len(text) > MAX_SEGMENT_CHARS:
            text = text[:MAX_SEGMENT_CHARS]
            result.truncated = True
        if total + len(text) > MAX_TOTAL_CHARS:
            text = text[: max(0, MAX_TOTAL_CHARS - total)]
            result.truncated = True
        total += len(text)
        seg.text = text
        capped.append(seg)
        if total >= MAX_TOTAL_CHARS:
            break
    result.segments = capped
    return result
